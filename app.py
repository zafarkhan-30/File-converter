from flask import Flask, render_template, request, send_from_directory
from docx import Document
import docx
import xlsxwriter
import os

app = Flask(__name__)

@app.route('/')
def upload_file():
    return render_template('upload.html')


@app.route('/uploader', methods=['GET', 'POST'])
def upload_file1():
    if request.method == 'POST':
        Questions_File = request.files['Questions_File']
        Answer_File = request.files['Answer_File']
        Questions_Doc = Document(Questions_File)
        Answer_Doc = Document(Answer_File)
        # print(Answer_Doc)

        name1 = os.path.splitext(Questions_File.filename)[0]
    
        Questions = []
        for paragraph in Questions_Doc.paragraphs:
            Questions.append(paragraph.text)

        # Answer_Doc
        Answer = []
        # print(Answer)
        for paragraph1 in Answer_Doc.paragraphs:
            Answer.append(paragraph1.text)

        Ans = []
        for character1 in Answer:
            rep = character1.replace("  ", "")
            Ans.append(rep)
            if '' in Ans:
                Ans.remove('')

        # print(Ans)
        Ans2 = []
        for i in Ans:
            if i == "(a)":
                Ans2.append(i)
            elif i == "(b)":
                Ans2.append(i)
            elif i == "(c)":
                Ans2.append(i)
            elif i == "(d)":
                Ans2.append(i)
        # print(Ans2)


        Que = []
        for character in Questions:
            d = character.replace("  ", "")
            c = d.replace("\t", "")
            Que.append(c)
            if '' in Que:
                Que.remove('')

        final_Questions = []
        
        for i in Que:
            x = (i.split("("))
            final_Questions.append(x)
            if '' in final_Questions:
                final_Questions.remove('')
        # print(final_Questions)

        dic = {}
        for i in final_Questions:
            for j in i:
                if j.startswith("a)"):
                    z = j.replace("a)", "")
                    if "option-A" not in dic:
                        dic["option-A"] = []
                    dic["option-A"].append(z)
                    # print(dic["option-A"])

                elif j.startswith("b)") or j.startswith("b "):
                    z = j.replace("b)", "") or j.replace("b", "")
                    if "option-B" not in dic:
                        dic["option-B"] = []
                    dic["option-B"].append(z)

                elif j.startswith("c)"):
                    z = j.replace("c)", "")
                    if "option-C" not in dic:
                        dic["option-C"] = []
                    dic["option-C"].append(z)

                elif j.startswith("d)") or j.startswith(("d")):
                    z = j.replace("d)", "")
                    if "option-D" not in dic:
                        dic["option-D"] = []
                    dic["option-D"].append(z)

                elif j.startswith("e)"):
                    z = j.replace("e)", "")
                    if "option-E" not in dic:
                        dic["option-E"] = []
                    dic["option-E"].append(z)
                else:
                    if "Questions" not in dic:
                        dic["Questions"] = []
                    dic["Questions"].append(j)

                    k = dic.values()
                    for x in k:
                        for z in x:
                            if z == '':
                                x.remove(z)
       
        for i in range(len(Ans2)):
            if Ans2[i] == "(a)":
                Ans2[i] = 1
            elif Ans2[i] == "(b)":
                Ans2[i] = 2
            elif Ans2[i] == "(c)":
                Ans2[i] = 3
            elif Ans2[i] == "(d)":
                Ans2[i] = 4
         

        for a in Ans2:
            if "Answer" not in dic:
                dic["Answer"] = []
            dic["Answer"].append(a)

            # if "Solution" not in dic:
            #     dic["Solution"] =[]
            # dic["Solution"].append(a)


        data = [dic]
        if name1.endswith("1"):
            workbook = xlsxwriter.Workbook("WithOutImageQuestion 1.xlsx")
            worksheet = workbook.add_worksheet("firstsheet")

            worksheet.write(0, 0, "Level")
            worksheet.write(0, 1, "Questions")
            worksheet.write(0, 2, "option-A")
            worksheet.write(0, 3, "option-B")
            worksheet.write(0, 4, "option-C")
            worksheet.write(0, 5, "option-D")
            # worksheet.write(0, 6, "option-E")
            worksheet.write(0, 7, "Answer")
            # worksheet.write(0,8 , "Solution")

            for index, entry in enumerate(data):
                content = []
                for i in range(200):
                    i = "M"
                    content.append(i)
                worksheet.write_column(index+1, 0, content)
                worksheet.write_column(index+1, 1, entry["Questions"])
                worksheet.write_column(index+1, 2, entry['option-A'])
                worksheet.write_column(index+1, 3, entry['option-B'])
                worksheet.write_column(index+1, 4, entry['option-C'])
                worksheet.write_column(index+1, 5, entry['option-D'])
                # worksheet.write_column(index+1, 6, entry['option-E'])
                worksheet.write_column(index+1, 7, entry['Answer'])
                # worksheet.write_column(index+1, 8, entry['Solution'])

            workbook.close()
            return send_from_directory("", "WithOutImageQuestion 1.xlsx", download_name="Question Answer sheet.xlsx")
        else:
            workbook = xlsxwriter.Workbook("WithOutImageQuestion 2.xlsx")
            worksheet = workbook.add_worksheet("firstsheet")

            worksheet.write(0, 0, "Level")
            worksheet.write(0, 1, "Questions")
            worksheet.write(0, 2, "option-A")
            worksheet.write(0, 3, "option-B")
            worksheet.write(0, 4, "option-C")
            worksheet.write(0, 5, "option-D")
            # worksheet.write(0, 6, "option-E")
            worksheet.write(0, 7, "Answer")
            # worksheet.write(0,8 , "Solution")

            for index, entry in enumerate(data):
                content = []
                for i in range(200):
                    i = "M"
                    content.append(i)

                worksheet.write_column(index+1, 0, content)
                worksheet.write_column(index+1, 1, entry["Questions"])
                worksheet.write_column(index+1, 2, entry['option-A'])
                worksheet.write_column(index+1, 3, entry['option-B'])
                worksheet.write_column(index+1, 4, entry['option-C'])
                worksheet.write_column(index+1, 5, entry['option-D'])
                # worksheet.write_column(index+1, 6, entry['option-E'])
                worksheet.write_column(index+1, 7, entry['Answer'])
                # worksheet.write_column(index+1, 8, entry['Solution'])

            workbook.close()
            return send_from_directory("", "WithOutImageQuestion 2.xlsx", download_name="Question Answer sheet 2.xlsx")


if __name__ == '__main__':
    app.run(debug=True)
